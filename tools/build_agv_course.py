from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path("/private/tmp/agv_course_output")
OUT.mkdir(parents=True, exist_ok=True)
STEM = "六年级_坐标仓储AGV_自动取货与返回_教学设计_20260623"

md = r"""# 六年级高阶项目教学设计

## 《坐标仓储AGV：自动取货与返回》

## 一、课程定位

- 适合年级：六年级高阶机器人课程
- 建议课时：连续2课时，每课时120分钟
- 课程类型：高阶算法综合项目
- 前置课程：直行算法、转向算法、函数封装、变量控制
- 项目重点：坐标导航、路线拆解、函数调用、自动取货与返回、误差校正

本项目安排在学生学习完直行算法、转向算法、函数封装和变量控制之后。学生不再只编写一段固定路线，而是通过输入目标坐标，让机器人读取数据、计算位置差值、调用基础运动函数，自动到达目标位置，完成取货后返回起点。

本项目不是循线任务，而是坐标导航任务。场地中的网格用于表示位置和距离，机器人移动依据来自目标坐标与当前位置之间的差值，不依赖黑线引导。

### 与四年级课程的区分

| 课程 | 寻找货架的方式 | 主要依据 | 学习重点 |
|---|---|---|---|
| 四年级《循线仓储叉车》 | 循线找货架 | 黑线、路口计数、传感器判断 | 机器人沿着环境走 |
| 六年级《坐标仓储AGV》 | 坐标找货架 | 输入坐标、计算x/y差值、执行直行和转向 | 机器人根据数据自己规划移动 |

四年级解决的是：机器人怎么沿着线找到货架。  
六年级解决的是：机器人怎么根据坐标自己找到位置。  
麦克纳姆轮解决的是：机器人怎么更高效地移动到位置。

因此，本项目既是四年级《循线仓储叉车》的算法升级，也是后续学习自动导航、路径优化和移动机器人控制的基础项目。

## 二、项目任务

### 基础任务

1. 机器人从起点 `(0,0)` 出发。
2. 教师或学生输入目标坐标，例如 `(2,1)`、`(3,2)`。
3. 机器人根据目标坐标自动移动到对应位置。
4. 到达目标位置后，完成取货动作。
5. 取货完成后，机器人自动返回起点。
6. 回到起点后完成卸货动作。

### 基础任务验收标准

- 更换目标坐标后，不需要重新编写整段移动程序。
- 机器人能够按照统一的“一格距离”移动到目标网格。
- 机器人到达目标位置后再执行取货动作。
- 取货完成后能够返回 `(0,0)`。
- 返回起点后执行卸货并停止。

### 进阶任务

1. 支持不同货架编号，例如A1、B2、C3。
2. 将货架编号转换为坐标。
3. 支持不同高度货架的取货动作。
4. 加入误差校正点。
5. 使用麦克纳姆轮实现横移或更高效路径。

## 三、课程目标

### 1. 知识目标

- 理解二维坐标系，能说出原点、x轴和y轴的含义。
- 理解当前位置、目标位置、x坐标和y坐标。
- 理解“一格距离”与机器人实际运动距离之间的关系。
- 理解目标坐标与当前位置之间的坐标差值。
- 理解普通双轮差速车需要通过转向改变运动方向。
- 理解麦克纳姆轮车可以横移，但仍需要准确的距离控制和误差校正。

### 2. 能力目标

- 编写并反复测试稳定的“直行一格”程序。
- 编写并测试准确的90度转向程序。
- 使用变量记录目标坐标和当前位置。
- 使用函数封装直行、转向、取货、卸货和返回动作。
- 根据目标坐标规划基础移动路线。
- 根据实际停车位置调试机器人移动误差。

### 3. 思维目标

- 建立用坐标描述位置的坐标化思维。
- 建立先拆分路线、再组合动作的路径规划意识。
- 建立重复功能只编写一次的函数封装意识。
- 建立测量、记录、比较和修正的误差校正意识。
- 能把综合任务拆分为定位、移动、取货、返回和卸货等子任务。

## 四、核心知识点

### 1. 坐标系统

以起点为原点 `(0,0)`，向右为x轴正方向，向前为y轴正方向。每一个方格代表机器人需要移动的固定距离。

```text
y
3    □    □    □    □
2    □    □    □    □
1    □    □    □    □
0   起点  □    □    □
     0    1    2    3   x
```

坐标 `(2,1)` 表示：从原点向x轴正方向移动2格，再向y轴正方向移动1格。课堂中应统一原点、坐标方向、车头初始朝向和每格长度。

### 2. 一格距离标定

坐标只是抽象位置，机器人最终执行的是电机运动。学生需要先完成“直行一格”函数标定，使输入1格、2格、3格时，实际位移能够保持基本稳定。

可采用以下关系理解程序：

```text
实际移动距离 = 格数 × 每格距离
电机执行量 = 格数 × 一格对应的电机参数
```

参数应通过同一起点的多次测试确定，不能只测试一次。

### 3. 坐标差值

机器人从当前位置 `(current_x, current_y)` 前往目标位置 `(target_x, target_y)` 时，需要计算：

```text
dx = target_x - current_x
dy = target_y - current_y
```

- `dx`决定x方向需要移动多少格。
- `dy`决定y方向需要移动多少格。
- 差值为正数时向正方向移动。
- 差值为负数时向反方向移动。
- 差值为0时，该方向不需要移动。

基础任务从 `(0,0)` 出发时，`dx = target_x`，`dy = target_y`。返回起点时，目标位置重新设为 `(0,0)`，或按去程动作的反向顺序执行。

### 4. 基础路线规划

普通双轮差速车可采用“先走x方向，再转向走y方向”的固定策略：

```text
读取目标坐标
计算dx、dy
沿x方向移动|dx|格
转向
沿y方向移动|dy|格
执行取货
按反向路线返回
执行卸货
```

基础阶段重点不是寻找最短路径，而是让路线规划规则清楚、可重复、可调试。

### 5. 函数封装

建议将重复动作封装为独立函数：

- `move_one_cell()`：直行一格。
- `move_cells(n)`：连续移动n格。
- `turn_left_90()`：左转90度。
- `turn_right_90()`：右转90度。
- `pick_up()`：执行取货。
- `unload()`：执行卸货。
- `go_to(x, y)`：根据坐标移动到目标位置。
- `return_home()`：返回原点。

函数名称可根据实际编程环境调整。课堂重点是让学生理解“主程序负责安排任务，函数负责完成具体动作”。

### 6. 普通双轮差速车与麦克纳姆轮车

| 对比项 | 普通双轮差速车 | 麦克纳姆轮车 |
|---|---|---|
| 横向移动 | 需要先转向再前进 | 可以直接横移 |
| 路线特点 | 动作清楚，适合学习转向和路线分解 | 路径更灵活，可减少部分转向 |
| 控制难点 | 90度转向精度、累计误差 | 多电机配合、横移偏差、方向组合 |
| 本项目定位 | 基础方案 | 进阶方案 |

麦克纳姆轮不会替代坐标计算。它改变的是“如何移动”，坐标系统解决的是“移动到哪里”。

### 7. 误差累积与校正

轮胎打滑、地面差异、车体重心和多次转向都会造成累计误差。调试时应先保证单个基础动作稳定，再组合完整任务。进阶任务可以在固定位置设置校正点，让机器人完成阶段性位置修正。

## 五、课堂流程

### 第一课时：建立坐标导航（120分钟）

#### 1. 项目导入与课程对比（15分钟）

- 回顾四年级循线仓储叉车如何寻找货架。
- 提出新问题：没有黑线时，机器人能否根据货架坐标找到位置？
- 明确坐标导航与循线导航的区别。

#### 2. 建立仓储坐标场地（15分钟）

- 确定原点 `(0,0)`、x轴、y轴和车头初始朝向。
- 统一每个网格的实际长度。
- 让学生根据坐标指出 `(1,2)`、`(3,1)` 等位置。

#### 3. 标定直行一格（25分钟）

- 调用已有直行算法，测试机器人移动一格。
- 从相同起点连续测试三次。
- 记录偏差，调整参数。
- 再测试连续移动2格和3格。

#### 4. 标定90度转向（20分钟）

- 测试左转和右转90度。
- 检查转向后车体是否与网格线方向保持一致。
- 分清转向误差与直行误差。

#### 5. 封装基础运动函数（20分钟）

- 封装直行一格、连续移动、左转和右转函数。
- 在主程序中只保留函数调用。
- 通过短路线检查各函数能否独立工作。

#### 6. 完成单目标坐标导航（20分钟）

- 输入一个目标坐标。
- 采用“先x后y”的规则拆分路线。
- 调用函数到达目标位置。
- 更换坐标，检查程序是否仍能运行。

#### 7. 课堂总结（5分钟）

- 学生说明坐标如何转换为移动格数。
- 保存已经通过测试的基础函数和参数。

### 第二课时：自动取货、返回与优化（120分钟）

#### 1. 复测基础函数（10分钟）

- 统一起点和车头方向。
- 快速复测直行一格和90度转向。
- 基础动作不稳定时先校准，不直接运行完整任务。

#### 2. 加入变量与坐标差值（20分钟）

- 使用变量保存目标坐标和当前位置。
- 计算 `dx` 与 `dy`。
- 根据差值决定移动方向和格数。

#### 3. 加入取货与卸货函数（20分钟）

- 将已有取货动作封装为 `pick_up()`。
- 将卸货动作封装为 `unload()`。
- 单独测试取货与卸货，确认动作完成后机构状态稳定。

#### 4. 编写返回起点逻辑（25分钟）

- 方案一：按去程路线的反向顺序返回。
- 方案二：把起点 `(0,0)` 作为新目标，再次调用坐标移动逻辑。
- 比较两种方案的程序结构和适用情况。

#### 5. 完整任务联调（30分钟）

- 输入目标坐标。
- 自动前往目标位置。
- 执行取货。
- 自动返回 `(0,0)`。
- 执行卸货。
- 每次只修改一个基础参数或一个程序环节。

#### 6. 多坐标验收（10分钟）

- 至少选择两个不同坐标进行测试。
- 记录到达位置、返回位置和主要偏差。
- 说明一次有效的调整。

#### 7. 总结与拓展说明（5分钟）

- 总结“输入数据—计算差值—调用函数—完成任务”的程序结构。
- 说明货架编号、校正点和麦克纳姆轮属于后续进阶方向。

## 六、学生常见问题

- 把坐标 `(x,y)` 的顺序写反。
- 没有统一原点、车头方向或每格长度。
- 只测试一格，却直接运行三格或四格路线。
- 把“电机转动量”直接当作“机器人移动距离”。
- 90度转向存在偏差，后续直行误差越来越大。
- 在函数内部和主程序中重复执行同一动作。
- 更新目标坐标后，没有同步更新当前位置。
- 返回时只把数值改成负数，却没有考虑车头朝向。
- 取货机构尚未复位就开始返回。
- 完整任务失败时同时修改多个参数，无法定位问题。

## 七、教师提醒

- 本课重点是坐标导航算法，不要把大量课堂时间消耗在重新设计取货机构上。
- 先确定坐标规则，再开始编程；原点、轴方向、初始朝向和每格长度必须全班统一。
- 直行一格和90度转向是完整项目的底层函数，必须先单独验收。
- 更换坐标时应只修改输入数据，不应重新复制一整段路线程序。
- 调试顺序建议为：单动作、短路线、单程取货、返回、完整任务。
- 普通双轮差速车与麦克纳姆轮车不要同时作为基础要求；先完成普通底盘方案，再开展进阶方案。
- 不直接提供最终运动参数，让学生依据记录的偏差进行调整。
- 课前确认端口连接，并确认取货机构动作不会与车体运动互相干涉。

## 八、拓展挑战

- 建立A1、B2、C3等货架编号与坐标的对应表。
- 输入货架编号后，自动转换为目标坐标。
- 使用变量保存当前朝向，减少不必要的转向。
- 比较“先走x再走y”和“先走y再走x”的运行差异。
- 增加中途校正点，降低长路线的累计误差。
- 设置多个取货任务，按顺序访问不同坐标。
- 为不同高度货架调用不同的取货函数。
- 使用麦克纳姆轮实现x方向横移和y方向直行，并比较任务效率。
- 在具备相应定位条件时，尝试闭环位置校正。

## 九、新老师上课提醒

- 课前准备：网格场地、坐标标签、已有机器人底盘、已有取货机构、测量工具和参数记录表。
- 本项目以前置算法为基础。学生如果还不能稳定完成直行一格和90度转向，应先补做基础动作测试。
- 第一课时必须完成坐标规则统一和基础函数标定，不要急着加入取货动作。
- 第二课时先复测基础函数，再加入坐标变量、取货和返回。
- 程序出现偏差时，先判断属于直行、转向、坐标计算、取货还是返回环节。
- 普通双轮差速车先使用固定的“先x后y”策略，避免一开始就引入过多路径选择。
- 麦克纳姆轮属于进阶方案，重点讲横移带来的路径变化，同时提醒学生仍需处理距离误差。
- 至少保留30分钟进行两个不同目标坐标的完整联调和验收。
- 课前确认端口连接；不确定的端口不在教案中固定编号，以现场机器人配置为准。
"""


def set_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.78)
section.right_margin = Inches(0.78)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.22

for style_name, size, color, before, after in [
    ("Heading 1", 15.5, (31, 78, 121), 14, 7),
    ("Heading 2", 12.5, (46, 116, 181), 10, 5),
    ("Heading 3", 11.3, (31, 78, 121), 7, 3),
]:
    st = styles[style_name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(*color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp.add_run("爱高创客｜六年级高阶机器人课程")
set_font(r, 8.5, True, (89, 102, 115))

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("坐标仓储AGV：自动取货与返回｜教师教学设计")
set_font(r, 8, False, (110, 120, 130))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(3)
r = p.add_run("六年级高阶项目教学设计")
set_font(r, 13, True, (46, 116, 181))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("《坐标仓储AGV：自动取货与返回》")
set_font(r, 23, True, (11, 37, 69))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
r = p.add_run("六年级高阶机器人课程｜连续2课时｜2026-06-23")
set_font(r, 9.5, False, (89, 102, 115))


def add_heading(text, level=1):
    doc.add_heading(text, level=level)


def add_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_font(r)


def add_numbers(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_font(r)


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F4F7")
    p_pr.append(shd)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, 9.2, False, (32, 55, 72), name="Menlo")


def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    for i, (cell, header_text, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        shade_cell(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header_text)
        set_font(r, 9, True, (31, 78, 121))
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for cell, value, width in zip(cells, row_data, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, 9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


add_heading("一、课程定位")
add_bullets([
    "适合年级：六年级高阶机器人课程",
    "建议课时：连续2课时，每课时120分钟",
    "课程类型：高阶算法综合项目",
    "前置课程：直行算法、转向算法、函数封装、变量控制",
    "项目重点：坐标导航、路线拆解、函数调用、自动取货与返回、误差校正",
])
add_para("本项目安排在学生学习完直行算法、转向算法、函数封装和变量控制之后。学生不再只编写一段固定路线，而是通过输入目标坐标，让机器人读取数据、计算位置差值、调用基础运动函数，自动到达目标位置，完成取货后返回起点。")
add_para("本项目不是循线任务，而是坐标导航任务。场地中的网格用于表示位置和距离，机器人移动依据来自目标坐标与当前位置之间的差值，不依赖黑线引导。")
add_heading("与四年级课程的区分", 2)
add_table(
    ["课程", "寻找货架的方式", "主要依据", "学习重点"],
    [
        ["四年级《循线仓储叉车》", "循线找货架", "黑线、路口计数、传感器判断", "机器人沿着环境走"],
        ["六年级《坐标仓储AGV》", "坐标找货架", "输入坐标、计算x/y差值、执行直行和转向", "机器人根据数据自己规划移动"],
    ],
    [1.55, 1.2, 2.45, 1.3],
)
add_para("四年级解决的是：机器人怎么沿着线找到货架。")
add_para("六年级解决的是：机器人怎么根据坐标自己找到位置。")
add_para("麦克纳姆轮解决的是：机器人怎么更高效地移动到位置。")
add_para("因此，本项目既是四年级《循线仓储叉车》的算法升级，也是后续学习自动导航、路径优化和移动机器人控制的基础项目。")

add_heading("二、项目任务")
add_heading("基础任务", 2)
add_numbers([
    "机器人从起点 (0,0) 出发。",
    "教师或学生输入目标坐标，例如 (2,1)、(3,2)。",
    "机器人根据目标坐标自动移动到对应位置。",
    "到达目标位置后，完成取货动作。",
    "取货完成后，机器人自动返回起点。",
    "回到起点后完成卸货动作。",
])
add_heading("基础任务验收标准", 2)
add_bullets([
    "更换目标坐标后，不需要重新编写整段移动程序。",
    "机器人能够按照统一的“一格距离”移动到目标网格。",
    "机器人到达目标位置后再执行取货动作。",
    "取货完成后能够返回 (0,0)。",
    "返回起点后执行卸货并停止。",
])
add_heading("进阶任务", 2)
add_numbers([
    "支持不同货架编号，例如A1、B2、C3。",
    "将货架编号转换为坐标。",
    "支持不同高度货架的取货动作。",
    "加入误差校正点。",
    "使用麦克纳姆轮实现横移或更高效路径。",
])

add_heading("三、课程目标")
add_heading("1. 知识目标", 2)
add_bullets([
    "理解二维坐标系，能说出原点、x轴和y轴的含义。",
    "理解当前位置、目标位置、x坐标和y坐标。",
    "理解“一格距离”与机器人实际运动距离之间的关系。",
    "理解目标坐标与当前位置之间的坐标差值。",
    "理解普通双轮差速车需要通过转向改变运动方向。",
    "理解麦克纳姆轮车可以横移，但仍需要准确的距离控制和误差校正。",
])
add_heading("2. 能力目标", 2)
add_bullets([
    "编写并反复测试稳定的“直行一格”程序。",
    "编写并测试准确的90度转向程序。",
    "使用变量记录目标坐标和当前位置。",
    "使用函数封装直行、转向、取货、卸货和返回动作。",
    "根据目标坐标规划基础移动路线。",
    "根据实际停车位置调试机器人移动误差。",
])
add_heading("3. 思维目标", 2)
add_bullets([
    "建立用坐标描述位置的坐标化思维。",
    "建立先拆分路线、再组合动作的路径规划意识。",
    "建立重复功能只编写一次的函数封装意识。",
    "建立测量、记录、比较和修正的误差校正意识。",
    "能把综合任务拆分为定位、移动、取货、返回和卸货等子任务。",
])

add_heading("四、核心知识点")
add_heading("1. 坐标系统", 2)
add_para("以起点为原点 (0,0)，向右为x轴正方向，向前为y轴正方向。每一个方格代表机器人需要移动的固定距离。")
add_code("y\n3    □    □    □    □\n2    □    □    □    □\n1    □    □    □    □\n0   起点  □    □    □\n     0    1    2    3   x")
add_para("坐标 (2,1) 表示：从原点向x轴正方向移动2格，再向y轴正方向移动1格。课堂中应统一原点、坐标方向、车头初始朝向和每格长度。")
add_heading("2. 一格距离标定", 2)
add_para("坐标只是抽象位置，机器人最终执行的是电机运动。学生需要先完成“直行一格”函数标定，使输入1格、2格、3格时，实际位移能够保持基本稳定。")
add_code("实际移动距离 = 格数 × 每格距离\n电机执行量 = 格数 × 一格对应的电机参数")
add_para("参数应通过同一起点的多次测试确定，不能只测试一次。")
add_heading("3. 坐标差值", 2)
add_para("机器人从当前位置 (current_x, current_y) 前往目标位置 (target_x, target_y) 时，需要计算：")
add_code("dx = target_x - current_x\ndy = target_y - current_y")
add_bullets([
    "dx决定x方向需要移动多少格。",
    "dy决定y方向需要移动多少格。",
    "差值为正数时向正方向移动。",
    "差值为负数时向反方向移动。",
    "差值为0时，该方向不需要移动。",
])
add_para("基础任务从 (0,0) 出发时，dx = target_x，dy = target_y。返回起点时，目标位置重新设为 (0,0)，或按去程动作的反向顺序执行。")
add_heading("4. 基础路线规划", 2)
add_para("普通双轮差速车可采用“先走x方向，再转向走y方向”的固定策略：")
add_code("读取目标坐标\n计算dx、dy\n沿x方向移动|dx|格\n转向\n沿y方向移动|dy|格\n执行取货\n按反向路线返回\n执行卸货")
add_para("基础阶段重点不是寻找最短路径，而是让路线规划规则清楚、可重复、可调试。")
add_heading("5. 函数封装", 2)
add_para("建议将重复动作封装为独立函数：")
add_bullets([
    "move_one_cell()：直行一格。",
    "move_cells(n)：连续移动n格。",
    "turn_left_90()：左转90度。",
    "turn_right_90()：右转90度。",
    "pick_up()：执行取货。",
    "unload()：执行卸货。",
    "go_to(x, y)：根据坐标移动到目标位置。",
    "return_home()：返回原点。",
])
add_para("函数名称可根据实际编程环境调整。课堂重点是让学生理解“主程序负责安排任务，函数负责完成具体动作”。")
add_heading("6. 普通双轮差速车与麦克纳姆轮车", 2)
add_table(
    ["对比项", "普通双轮差速车", "麦克纳姆轮车"],
    [
        ["横向移动", "需要先转向再前进", "可以直接横移"],
        ["路线特点", "动作清楚，适合学习转向和路线分解", "路径更灵活，可减少部分转向"],
        ["控制难点", "90度转向精度、累计误差", "多电机配合、横移偏差、方向组合"],
        ["本项目定位", "基础方案", "进阶方案"],
    ],
    [1.2, 2.65, 2.65],
)
add_para("麦克纳姆轮不会替代坐标计算。它改变的是“如何移动”，坐标系统解决的是“移动到哪里”。")
add_heading("7. 误差累积与校正", 2)
add_para("轮胎打滑、地面差异、车体重心和多次转向都会造成累计误差。调试时应先保证单个基础动作稳定，再组合完整任务。进阶任务可以在固定位置设置校正点，让机器人完成阶段性位置修正。")

add_heading("五、课堂流程")
add_heading("第一课时：建立坐标导航（120分钟）", 2)
flow1 = [
    ("1. 项目导入与课程对比（15分钟）", ["回顾四年级循线仓储叉车如何寻找货架。", "提出新问题：没有黑线时，机器人能否根据货架坐标找到位置？", "明确坐标导航与循线导航的区别。"]),
    ("2. 建立仓储坐标场地（15分钟）", ["确定原点 (0,0)、x轴、y轴和车头初始朝向。", "统一每个网格的实际长度。", "让学生根据坐标指出 (1,2)、(3,1) 等位置。"]),
    ("3. 标定直行一格（25分钟）", ["调用已有直行算法，测试机器人移动一格。", "从相同起点连续测试三次。", "记录偏差，调整参数。", "再测试连续移动2格和3格。"]),
    ("4. 标定90度转向（20分钟）", ["测试左转和右转90度。", "检查转向后车体是否与网格线方向保持一致。", "分清转向误差与直行误差。"]),
    ("5. 封装基础运动函数（20分钟）", ["封装直行一格、连续移动、左转和右转函数。", "在主程序中只保留函数调用。", "通过短路线检查各函数能否独立工作。"]),
    ("6. 完成单目标坐标导航（20分钟）", ["输入一个目标坐标。", "采用“先x后y”的规则拆分路线。", "调用函数到达目标位置。", "更换坐标，检查程序是否仍能运行。"]),
    ("7. 课堂总结（5分钟）", ["学生说明坐标如何转换为移动格数。", "保存已经通过测试的基础函数和参数。"]),
]
for title, items in flow1:
    add_heading(title, 3)
    add_bullets(items)
add_heading("第二课时：自动取货、返回与优化（120分钟）", 2)
flow2 = [
    ("1. 复测基础函数（10分钟）", ["统一起点和车头方向。", "快速复测直行一格和90度转向。", "基础动作不稳定时先校准，不直接运行完整任务。"]),
    ("2. 加入变量与坐标差值（20分钟）", ["使用变量保存目标坐标和当前位置。", "计算dx与dy。", "根据差值决定移动方向和格数。"]),
    ("3. 加入取货与卸货函数（20分钟）", ["将已有取货动作封装为pick_up()。", "将卸货动作封装为unload()。", "单独测试取货与卸货，确认动作完成后机构状态稳定。"]),
    ("4. 编写返回起点逻辑（25分钟）", ["方案一：按去程路线的反向顺序返回。", "方案二：把起点 (0,0) 作为新目标，再次调用坐标移动逻辑。", "比较两种方案的程序结构和适用情况。"]),
    ("5. 完整任务联调（30分钟）", ["输入目标坐标。", "自动前往目标位置。", "执行取货。", "自动返回 (0,0)。", "执行卸货。", "每次只修改一个基础参数或一个程序环节。"]),
    ("6. 多坐标验收（10分钟）", ["至少选择两个不同坐标进行测试。", "记录到达位置、返回位置和主要偏差。", "说明一次有效的调整。"]),
    ("7. 总结与拓展说明（5分钟）", ["总结“输入数据—计算差值—调用函数—完成任务”的程序结构。", "说明货架编号、校正点和麦克纳姆轮属于后续进阶方向。"]),
]
for title, items in flow2:
    add_heading(title, 3)
    add_bullets(items)

add_heading("六、学生常见问题")
add_bullets([
    "把坐标 (x,y) 的顺序写反。",
    "没有统一原点、车头方向或每格长度。",
    "只测试一格，却直接运行三格或四格路线。",
    "把“电机转动量”直接当作“机器人移动距离”。",
    "90度转向存在偏差，后续直行误差越来越大。",
    "在函数内部和主程序中重复执行同一动作。",
    "更新目标坐标后，没有同步更新当前位置。",
    "返回时只把数值改成负数，却没有考虑车头朝向。",
    "取货机构尚未复位就开始返回。",
    "完整任务失败时同时修改多个参数，无法定位问题。",
])

add_heading("七、教师提醒")
add_bullets([
    "本课重点是坐标导航算法，不要把大量课堂时间消耗在重新设计取货机构上。",
    "先确定坐标规则，再开始编程；原点、轴方向、初始朝向和每格长度必须全班统一。",
    "直行一格和90度转向是完整项目的底层函数，必须先单独验收。",
    "更换坐标时应只修改输入数据，不应重新复制一整段路线程序。",
    "调试顺序建议为：单动作、短路线、单程取货、返回、完整任务。",
    "普通双轮差速车与麦克纳姆轮车不要同时作为基础要求；先完成普通底盘方案，再开展进阶方案。",
    "不直接提供最终运动参数，让学生依据记录的偏差进行调整。",
    "课前确认端口连接，并确认取货机构动作不会与车体运动互相干涉。",
])

add_heading("八、拓展挑战")
add_bullets([
    "建立A1、B2、C3等货架编号与坐标的对应表。",
    "输入货架编号后，自动转换为目标坐标。",
    "使用变量保存当前朝向，减少不必要的转向。",
    "比较“先走x再走y”和“先走y再走x”的运行差异。",
    "增加中途校正点，降低长路线的累计误差。",
    "设置多个取货任务，按顺序访问不同坐标。",
    "为不同高度货架调用不同的取货函数。",
    "使用麦克纳姆轮实现x方向横移和y方向直行，并比较任务效率。",
    "在具备相应定位条件时，尝试闭环位置校正。",
])

doc.add_page_break()
add_heading("九、新老师上课提醒")
add_bullets([
    "课前准备：网格场地、坐标标签、已有机器人底盘、已有取货机构、测量工具和参数记录表。",
    "本项目以前置算法为基础。学生如果还不能稳定完成直行一格和90度转向，应先补做基础动作测试。",
    "第一课时必须完成坐标规则统一和基础函数标定，不要急着加入取货动作。",
    "第二课时先复测基础函数，再加入坐标变量、取货和返回。",
    "程序出现偏差时，先判断属于直行、转向、坐标计算、取货还是返回环节。",
    "普通双轮差速车先使用固定的“先x后y”策略，避免一开始就引入过多路径选择。",
    "麦克纳姆轮属于进阶方案，重点讲横移带来的路径变化，同时提醒学生仍需处理距离误差。",
    "至少保留30分钟进行两个不同目标坐标的完整联调和验收。",
    "课前确认端口连接；不确定的端口不在教案中固定编号，以现场机器人配置为准。",
])

md_path = OUT / f"{STEM}.md"
docx_path = OUT / f"{STEM}.docx"
md_path.write_text(md, encoding="utf-8")
doc.save(docx_path)
print(md_path)
print(docx_path)
